#!/usr/bin/env python3
"""
Script de conversion du plan de présentation MD vers DOCX
"""

import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Couleurs Argam
GOLDEN_BROWN = RGBColor(0xB4, 0x92, 0x5E)
AUBERGINE = RGBColor(0x52, 0x4C, 0x5D)


def create_docx_from_md():
    # Lire le fichier MD
    with open(
        "/Users/julianbonne/Desktop/ARGAMCONSEIL NEW/PRESENTATION_COGOHR_PLAN.md",
        "r",
        encoding="utf-8",
    ) as f:
        content = f.read()

    # Créer le document
    doc = Document()

    # Configurer les styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Parcourir le contenu ligne par ligne
    lines = content.split("\n")
    in_table = False
    table_rows = []
    in_script = False

    for line in lines:
        line = line.rstrip()

        # Ignorer les lignes de séparation
        if line == "---":
            doc.add_paragraph()
            continue

        # Tables
        if "|" in line and not line.startswith(">"):
            if "---" in line:
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(cells)
            continue
        elif in_table:
            # Fin de table, créer la table
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = "Table Grid"
                for i, row_data in enumerate(table_rows):
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        cell = row.cells[j]
                        cell.text = cell_text.replace("**", "")
                        if i == 0:  # Header row
                            cell.paragraphs[0].runs[0].bold = True
            in_table = False
            table_rows = []

        # Titre H1
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Titre H2
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue

        # Titre H3
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue

        # Titre H4
        if line.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Citations/Scripts
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(line[2:].strip('"'))
            run.italic = True
            run.font.color.rgb = AUBERGINE
            continue

        # Listes avec checkbox
        if line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ " + line[6:])
            continue

        # Listes à puces
        if line.startswith("- "):
            text = line[2:]
            # Traiter le gras
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
            continue

        # Listes numérotées
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line)
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
            continue

        # Texte normal avec formatage
        if line.strip():
            # Nettoyer le formatage markdown
            text = line
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Gras
            text = re.sub(r"\*(.+?)\*", r"\1", text)  # Italique

            p = doc.add_paragraph()

            # Détecter si c'est un texte en gras uniquement
            if line.startswith("**") and line.endswith("**"):
                run = p.add_run(text)
                run.bold = True
            else:
                p.add_run(text)

    # Sauvegarder
    output_path = (
        "/Users/julianbonne/Desktop/ARGAMCONSEIL NEW/PRESENTATION_COGOHR_PLAN.docx"
    )
    doc.save(output_path)
    print(f"✅ Document créé : {output_path}")


if __name__ == "__main__":
    create_docx_from_md()
