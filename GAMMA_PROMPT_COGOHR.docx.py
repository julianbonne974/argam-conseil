#!/usr/bin/env python3
"""
Création du fichier DOCX optimisé pour Gamma.app
Bonnes pratiques appliquées :
- Contexte clair (audience + objectif)
- Structure slide par slide
- Bullet points courts (30 mots max par slide)
- Pas de données inventées
- Ton et style définis
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def create_gamma_prompt():
    doc = Document()

    # Titre principal
    title = doc.add_heading("PROMPT GAMMA.APP - Présentation COGOHR", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section Context
    doc.add_heading("CONTEXTE & INSTRUCTIONS GAMMA", level=1)

    context = doc.add_paragraph()
    context.add_run("Audience : ").bold = True
    context.add_run(
        "Adhérents COGOHR (Comité des Oeuvres Sociales des Hôpitaux de La Réunion), fonctionnaires hospitaliers.\n\n"
    )

    context.add_run("Objectif : ").bold = True
    context.add_run(
        "Présenter le partenariat Argam Conseils x COGOHR et convaincre les adhérents de demander une étude gratuite pour leur retraite.\n\n"
    )

    context.add_run("Ton : ").bold = True
    context.add_run(
        "Professionnel, rassurant, accessible. Pas de jargon financier complexe.\n\n"
    )

    context.add_run("Style visuel : ").bold = True
    context.add_run(
        "Couleurs : doré (#B4925E) et aubergine (#524C5D). Épuré, moderne, corporate.\n\n"
    )

    context.add_run("Format : ").bold = True
    context.add_run(
        "10 slides maximum. 30 mots max par slide. Bullet points courts. Une idée par slide.\n\n"
    )

    context.add_run("Contrainte : ").bold = True
    context.add_run(
        "Ne pas inventer de statistiques. Utiliser uniquement les données fournies ci-dessous."
    )

    doc.add_page_break()

    # ===== SLIDES =====
    doc.add_heading("CONTENU DES 10 SLIDES", level=1)

    # Slide 1
    doc.add_heading("SLIDE 1 : Titre", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Argam Conseils x COGOHR\n")
    p.add_run("Sous-titre : ").bold = True
    p.add_run("Protégez votre niveau de vie à la retraite\n")
    p.add_run("Présenté par : ").bold = True
    p.add_run("Clémence Gavache\n")
    p.add_run("Badge : ").bold = True
    p.add_run("Partenariat Exclusif")

    # Slide 2
    doc.add_heading("SLIDE 2 : Qui sommes-nous ?", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Argam Conseils - Votre partenaire patrimoine\n\n")
    p.add_run("Bullet points :\n").bold = True
    doc.add_paragraph(
        "15+ années d'expertise en conseil patrimonial", style="List Bullet"
    )
    doc.add_paragraph("900+ clients fonctionnaires accompagnés", style="List Bullet")
    doc.add_paragraph(
        "Spécialistes de la fonction publique hospitalière", style="List Bullet"
    )
    doc.add_paragraph("Conseil indépendant et objectif", style="List Bullet")

    # Slide 3
    doc.add_heading("SLIDE 3 : Notre présence", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Une équipe proche de vous\n\n")
    p.add_run("Chiffres clés :\n").bold = True
    doc.add_paragraph("8 personnes à votre service", style="List Bullet")
    doc.add_paragraph("2 implantations : Bordeaux et La Réunion", style="List Bullet")
    doc.add_paragraph("Connaissance des spécificités locales", style="List Bullet")
    doc.add_paragraph("Rendez-vous en visio ou présentiel", style="List Bullet")

    # Slide 4
    doc.add_heading("SLIDE 4 : Le problème", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("53% de vos revenus ne comptent pas pour la retraite\n\n")
    p.add_run("Message clé :\n").bold = True
    p.add_run(
        "Votre prime de vie chère de 53% n'est PAS prise en compte dans le calcul de votre pension de retraite.\n\n"
    )
    p.add_run("Visuel suggéré : ").bold = True
    p.add_run("Graphique camembert 53% prime / 47% traitement indiciaire")

    # Slide 5
    doc.add_heading("SLIDE 5 : L'impact chiffré", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Votre situation à la retraite\n\n")
    p.add_run("Comparaison :\n").bold = True
    doc.add_paragraph("Aujourd'hui : 3 000 €/mois", style="List Bullet")
    doc.add_paragraph("À la retraite : 1 400 €/mois", style="List Bullet")
    doc.add_paragraph("Perte : jusqu'à -53% de revenus", style="List Bullet")
    p2 = doc.add_paragraph()
    p2.add_run("\nVisuel suggéré : ").bold = True
    p2.add_run("Deux boxes comparatives avec flèche descendante rouge")

    # Slide 6
    doc.add_heading("SLIDE 6 : La solution PER", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Le Plan Épargne Retraite\n\n")
    p.add_run("3 avantages :\n").bold = True
    doc.add_paragraph("Déduction fiscale immédiate sur vos impôts", style="List Bullet")
    doc.add_paragraph(
        "Complément de retraite : reconstituez jusqu'à 30% de vos revenus",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Sortie flexible : capital, rente ou les deux", style="List Bullet"
    )

    # Slide 7
    doc.add_heading("SLIDE 7 : Offre exclusive COGOHR (1/2)", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Étude patrimoniale GRATUITE\n\n")
    p.add_run("Valeur : 300€ - Offerte aux adhérents COGOHR\n\n").bold = True
    p.add_run("Inclus :\n").bold = True
    doc.add_paragraph("Analyse personnalisée de votre situation", style="List Bullet")
    doc.add_paragraph("Simulation retraite détaillée", style="List Bullet")
    doc.add_paragraph("Recommandations adaptées", style="List Bullet")
    doc.add_paragraph("Sans engagement - Durée : 45 minutes", style="List Bullet")

    # Slide 8
    doc.add_heading("SLIDE 8 : Offre exclusive COGOHR (2/2)", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("0% de frais d'entrée\n\n")
    p.add_run("Avantages :\n").bold = True
    doc.add_paragraph("Aucun frais de dossier", style="List Bullet")
    doc.add_paragraph("Aucun frais d'ouverture", style="List Bullet")
    doc.add_paragraph("100% de vos versements investis", style="List Bullet")
    doc.add_paragraph("Économie jusqu'à 3% de vos versements", style="List Bullet")
    p2 = doc.add_paragraph()
    p2.add_run("\nRésultat : ").bold = True
    p2.add_run("Plusieurs milliers d'euros économisés")

    # Slide 9
    doc.add_heading("SLIDE 9 : Notre processus", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("3 étapes simples\n\n")
    p.add_run("Étapes :\n").bold = True
    doc.add_paragraph(
        "1. DIAGNOSTIC : Analyse complète gratuite (45 min, visio ou présentiel)",
        style="List Bullet",
    )
    doc.add_paragraph(
        "2. SOLUTION : Plan sur-mesure avec simulation chiffrée", style="List Bullet"
    )
    doc.add_paragraph(
        "3. SUIVI : Accompagnement annuel avec conseiller dédié", style="List Bullet"
    )

    # Slide 10
    doc.add_heading("SLIDE 10 : Appel à l'action", level=2)
    p = doc.add_paragraph()
    p.add_run("Titre : ").bold = True
    p.add_run("Demandez votre étude gratuite\n\n")
    p.add_run("Contact :\n").bold = True
    doc.add_paragraph("Téléphone : 05 33 89 14 00", style="List Bullet")
    doc.add_paragraph("Email : contact-reunion@argamconseils.com", style="List Bullet")
    doc.add_paragraph("Réponse garantie sous 48h", style="List Bullet")
    p2 = doc.add_paragraph()
    p2.add_run("\nGaranties : ").bold = True
    p2.add_run("Sans engagement • 100% confidentiel • Réservé adhérents COGOHR")

    doc.add_page_break()

    # Instructions finales pour Gamma
    doc.add_heading("INSTRUCTIONS POUR GAMMA.APP", level=1)

    instructions = """
Copiez le prompt ci-dessous dans Gamma.app :

---

Crée une présentation professionnelle de 10 slides pour Argam Conseils, cabinet de conseil en gestion de patrimoine.

CONTEXTE :
- Audience : Fonctionnaires hospitaliers de La Réunion, adhérents COGOHR
- Objectif : Les convaincre de demander une étude retraite gratuite
- Ton : Professionnel, rassurant, accessible

STYLE VISUEL :
- Couleurs principales : doré (#B4925E) et aubergine (#524C5D)
- Style épuré et moderne
- Maximum 30 mots par slide
- Utiliser des icônes professionnelles

STRUCTURE DES 10 SLIDES :

1. TITRE : "Argam Conseils x COGOHR - Protégez votre niveau de vie à la retraite" (Présenté par Clémence Gavache, Badge: Partenariat Exclusif)

2. QUI SOMMES-NOUS : 15+ ans d'expertise, 900+ clients, spécialistes fonction publique hospitalière, conseil indépendant

3. NOTRE PRÉSENCE : 8 personnes, 2 implantations (Bordeaux + La Réunion), rendez-vous visio ou présentiel

4. LE PROBLÈME : 53% de prime de vie chère NON prise en compte pour la retraite (graphique camembert)

5. L'IMPACT : Aujourd'hui 3000€/mois → Retraite 1400€/mois = -53% de revenus

6. LA SOLUTION PER : Déduction fiscale immédiate + Jusqu'à 30% revenus reconstitués + Sortie flexible (capital/rente)

7. OFFRE COGOHR 1 : Étude patrimoniale GRATUITE (valeur 300€) - Analyse personnalisée, simulation détaillée, 45min, sans engagement

8. OFFRE COGOHR 2 : 0% frais d'entrée - 100% versements investis - Économie jusqu'à 3% = plusieurs milliers d'euros

9. PROCESSUS : 3 étapes (Diagnostic gratuit 45min → Solution sur-mesure → Suivi annuel avec conseiller dédié)

10. CTA : Demandez votre étude gratuite - Tél: 05 33 89 14 00 - Email: contact-reunion@argamconseils.com - Réponse sous 48h - Sans engagement

CONTRAINTES :
- Ne pas inventer de statistiques
- Utiliser uniquement les chiffres fournis
- Mettre en avant "GRATUIT" et "0%" visuellement
"""

    doc.add_paragraph(instructions)

    # Sauvegarder
    output_path = "/Users/julianbonne/Desktop/ARGAMCONSEIL NEW/GAMMA_PROMPT_COGOHR.docx"
    doc.save(output_path)
    print(f"✅ Document Gamma créé : {output_path}")


if __name__ == "__main__":
    create_gamma_prompt()
